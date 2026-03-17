from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from datetime import datetime
import os
import requests
import json
import argparse
import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

# ----------------------
# Fetch book info and format to MLA
# ----------------------
def fetch_book_mla(author, title):
    """Fetch book information and format to MLA citation."""
    publisher = 'Publisher'
    year = 'Year'
    authors = [author]
    title_full = title

    # Try Google Books API first
    try:
        url = f"https://www.googleapis.com/books/v1/volumes?q=intitle:{title}+inauthor:{author}&maxResults=1"
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json()
            if 'items' in data and data['items']:
                item = data['items'][0]['volumeInfo']
                authors = item.get('authors', authors)
                title_full = item.get('title', title_full)
                pub = item.get('publisher')
                if pub:
                    publisher = pub
                    print(f"Retrieved from Google Books: Publisher={publisher}")
                yr = item.get('publishedDate')
                if yr:
                    year = yr[:4]
                    print(f"Retrieved from Google Books: Year={year}")
                if publisher != 'Publisher':
                    print(f"Using Google Books data for {title}")
                    author_str = ', '.join([f"{a.split()[-1]}, {' '.join(a.split()[:-1])}" for a in authors if a])
                    return f"{author_str}. {title_full}. {publisher}, {year}."
    except Exception as e:
        print(f"Google Books failed: {e}")

    # If publisher still default, try Open Library
    if publisher == 'Publisher':
        try:
            url = f"https://openlibrary.org/search.json?title={title}&author={author}&limit=1"
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                data = response.json()
                if 'docs' in data and data['docs']:
                    doc = data['docs'][0]
                    auths = doc.get('author_name')
                    if auths:
                        authors = auths
                    tit = doc.get('title')
                    if tit:
                        title_full = tit
                    pubs = doc.get('publisher')
                    if pubs and pubs[0]:
                        publisher = pubs[0]
                        print(f"Retrieved from Open Library: Publisher={publisher}")
                    yr = doc.get('first_publish_year')
                    if yr:
                        year = str(yr)
                        print(f"Retrieved from Open Library: Year={year}")
                    if publisher != 'Publisher':
                        print(f"Using Open Library data for {title}")
                        author_str = ', '.join([f"{a.split()[-1]}, {' '.join(a.split()[:-1])}" for a in authors if a])
                        return f"{author_str}. {title_full}. {publisher}, {year}."
        except Exception as e:
            print(f"Open Library failed: {e}")

    # If still default, try CrossRef
    if publisher == 'Publisher':
        try:
            url = f"https://api.crossref.org/works?query.title={title}&query.author-name={author}&rows=1"
            response = requests.get(url, timeout=5)
            if response.status_code == 200:
                data = response.json()
                if 'message' in data and 'items' in data['message'] and data['message']['items']:
                    item = data['message']['items'][0]
                    auths = item.get('author', [])
                    if auths:
                        authors = [f"{a.get('given', '')} {a.get('family', '')}".strip() for a in auths if a.get('family')]
                    tit = item.get('title')
                    if tit and tit[0]:
                        title_full = tit[0]
                    pub = item.get('publisher')
                    if pub:
                        publisher = pub
                        print(f"Retrieved from CrossRef: Publisher={publisher}")
                    pub_date = item.get('published-print') or item.get('published-online')
                    if pub_date and 'date-parts' in pub_date and pub_date['date-parts'][0]:
                        year = str(pub_date['date-parts'][0][0])
                        print(f"Retrieved from CrossRef: Year={year}")
                    print(f"Using CrossRef data for {title}")
                    author_str = ', '.join([f"{a.split()[-1]}, {' '.join(a.split()[:-1])}" for a in authors if a])
                    return f"{author_str}. {title_full}. {publisher}, {year}."
        except Exception as e:
            print(f"CrossRef failed: {e}")

    # Fallback
    print(f"No data found, using fallback for {title}")
    author_str = ', '.join([f"{a.split()[-1]}, {' '.join(a.split()[:-1])}" for a in authors if a])
    return f"{author_str}. {title_full}. {publisher}, {year}."

# win32 (pywin32) is optional — used only for accurate page counts when available
try:
    import win32com.client as win32
    WIN32_AVAILABLE = True
except Exception:
    win32 = None
    WIN32_AVAILABLE = False

# ----------------------
# Word 페이지 수 확인
# ----------------------
def get_page_count(doc_path):
    """Return a page count for a .docx document.

    If win32com is available (Windows), use Word's accurate page count.
    Otherwise, fall back to a rough estimate based on text length.
    """
    # Ensure doc_path is absolute
    abs_doc_path = os.path.abspath(doc_path)

    if WIN32_AVAILABLE:
        # Initialize Word application FIRST, before opening documents
        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False

        # Open the document only once, using absolute path
        doc = word.Documents.Open(abs_doc_path)
        doc.Repaginate()
        pages = doc.ComputeStatistics(2)  # 2: wdStatisticPages

        doc.Close(False)
        word.Quit()

        return pages

    # Fallback estimation (cross-platform)
    doc = Document(abs_doc_path)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Assume roughly 1800 characters per page (approximation)
    approx_chars_per_page = 1800
    pages = max(1, (len(text) + approx_chars_per_page - 1) // approx_chars_per_page)
    return pages

def set_document_font(doc, font_name):
    """문서 전체 글꼴 변경"""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(12)

    for section in doc.sections:
        header = section.header
        for p in header.paragraphs:
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(12)


def _insert_paragraph_after(paragraph, text="", style=None):
    """Insert a new paragraph immediately after the given paragraph."""
    new_p_elm = OxmlElement('w:p')
    paragraph._p.addnext(new_p_elm)
    new_para = Paragraph(new_p_elm, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def _copy_run_format(src_run, dst_run):
    """Copy visual formatting from one run to another."""
    for attr in ("bold", "italic", "underline"):
        if hasattr(src_run, attr):
            setattr(dst_run, attr, getattr(src_run, attr))

    # Copy font properties if available
    if src_run.font is not None:
        if src_run.font.name is not None:
            dst_run.font.name = src_run.font.name
        if src_run.font.size is not None:
            dst_run.font.size = src_run.font.size
        if src_run.font.color is not None and src_run.font.color.rgb is not None:
            dst_run.font.color.rgb = src_run.font.color.rgb
        if hasattr(src_run.font, "highlight_color"):
            dst_run.font.highlight_color = src_run.font.highlight_color


def _split_paragraph_on_newlines(paragraph):
    """Split a paragraph into multiple paragraphs at newline characters.

    This preserves run-level formatting by copying run properties into the newly
    created paragraphs.
    """
    if "\n" not in paragraph.text:
        return

    style = paragraph.style
    # capture relevant paragraph format settings
    pformat = paragraph.paragraph_format
    format_values = {
        "alignment": pformat.alignment,
        "left_indent": pformat.left_indent,
        "right_indent": pformat.right_indent,
        "first_line_indent": pformat.first_line_indent,
        "space_before": pformat.space_before,
        "space_after": pformat.space_after,
        "line_spacing": pformat.line_spacing,
    }

    segments = [[]]
    for run in list(paragraph.runs):
        parts = run.text.split("\n")
        for idx, part in enumerate(parts):
            segments[-1].append((run, part))
            if idx < len(parts) - 1:
                segments.append([])

    # rebuild the paragraph and insert new paragraphs for extra segments
    paragraph._p.clear_content()
    paragraph.style = style
    for k, v in format_values.items():
        setattr(paragraph.paragraph_format, k, v)

    def _build_paragraph(par, segment):
        for src_run, text in segment:
            if not text:
                continue
            dst = par.add_run(text)
            _copy_run_format(src_run, dst)

    # first segment stays in the original paragraph
    _build_paragraph(paragraph, segments[0])

    current_paragraph = paragraph
    for segment in segments[1:]:
        current_paragraph = _insert_paragraph_after(current_paragraph, style=style)
        for k, v in format_values.items():
            setattr(current_paragraph.paragraph_format, k, v)
        _build_paragraph(current_paragraph, segment)

# ----------------------
# 페이지 번호 추가
# ----------------------
def add_page_number(header, name=None):
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"{name} " if name else "")

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    # Ensure header font consistent
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def apply_common_styles(doc):
    """Apply common Normal style settings to the document."""
    if 'Normal' in doc.styles:
        style = doc.styles['Normal']
        set_document_font(doc, "Cambria")
        font = style.font
        font.name = "Cambria"
        font.size = Pt(12)
        pformat = style.paragraph_format
        pformat.line_spacing = 2.0
        pformat.space_before = Pt(0)
        pformat.space_after = Pt(0)
        pformat.first_line_indent = Inches(0.5)

# ----------------------
# 문서 포맷팅
# ----------------------
def mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit=None, include_works_cited_in_limit=False, works_cited_input=None):

    doc = Document(input_file)

    # 섹션 마진 설정 (기본 MLA: 1-inch margins)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # 페이지 헤더 (MLA: 성 + 페이지)
    header = section.header
    last_name = student_name.split()[-1] if student_name else None
    add_page_number(header, last_name)

    # 공통 스타일 적용
    apply_common_styles(doc)

    # 문서 맨 앞 정보 삽입 (MLA 스타일 상단 왼쪽에 배치)
    header_info = [
        student_name,
        professor_name,
        course_name,
        datetime.today().strftime("%d %B %Y")
    ]
    for text in reversed(header_info):
        p = doc.paragraphs[0].insert_paragraph_before(text)
        p.paragraph_format.first_line_indent = Inches(0)

    # 제목 추가 (가운데 정렬)
    title_paragraph = doc.paragraphs[len(header_info)].insert_paragraph_before(title)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.first_line_indent = Inches(0)

    # Find the title paragraph
    title_index = next((i for i,p in enumerate(doc.paragraphs) if p.text == title), None)
    if title_index is None:
        title_index = len(header_info)  # fallback

    # Split all paragraphs on newlines
    for p in list(doc.paragraphs[title_index + 1:]):
        if '\n' in p.text:
            _split_paragraph_on_newlines(p)

    # Apply indentation after every \n (each new paragraph gets first-line indent if not already indented)
    for p in doc.paragraphs[title_index + 1:]:
        text = p.text
        if text.strip() == "":
            continue
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if text == text.lstrip():  # no leading whitespace, so add indentation
            p.paragraph_format.first_line_indent = Inches(0.5)
        # else, already has indentation, leave as is

    # Ensure a "Works Cited" section exists on the last page
    works_cited_added = False
    last_chunk = "\n".join(p.text for p in doc.paragraphs[-5:]).lower()
    if "works cited" not in last_chunk:
        # Add a page break then add the heading centered.
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        works_cited_heading = doc.add_paragraph("Works Cited")
        works_cited_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        works_cited_heading.paragraph_format.first_line_indent = Inches(0)
        works_cited_added = True

        # Add works cited entries if provided
        if works_cited_input:
            entries = works_cited_input.split('\n')
            for entry in entries:
                entry = entry.strip()
                if entry:
                    # Check if already in MLA format (contains period and comma)
                    if '.' in entry and ',' in entry:
                        # Assume MLA format
                        p = doc.add_paragraph(entry)
                    else:
                        # Assume author: title
                        if ':' in entry:
                            author, title = entry.split(':', 1)
                            author = author.strip()
                            title = title.strip()
                            mla_entry = fetch_book_mla(author, title)
                            p = doc.add_paragraph(mla_entry)
                        else:
                            # Fallback
                            p = doc.add_paragraph(entry)
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 페이지 제한 체크
    if page_limit is not None:

        doc.save(output_file)

        pages = get_page_count(output_file)

        # Adjust page count if Works Cited is not to be included in limit
        if not include_works_cited_in_limit and works_cited_added:
            pages -= 1

        if pages > page_limit:
            print(f"Page limit exceeded ({pages} > {page_limit})")
            print("Switching font from Cambria to Times New Roman...")

            # Switch the whole document to Times New Roman to try to reduce page count.
            set_document_font(doc, "Times New Roman")
            if 'Normal' in doc.styles:
                doc.styles['Normal'].font.name = "Times New Roman"

            doc.save(output_file)
            pages = get_page_count(output_file)

            if pages > page_limit:
                print("⚠ WARNING: Even with Times New Roman the document exceeds the page limit.")
            else:
                print("Page count fixed by switching to Times New Roman.")
        else:
            print(f"Pages within limit ({pages} <= {page_limit})")

    # 최종 저장
    doc.save(output_file)
    print(f"MLA-formatted document saved as {output_file}")

# ----------------------
# Other format stubs (not implemented)
# ----------------------
def chicago_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit=None, include_works_cited_in_limit=False, works_cited_input=None):
    print("Chicago format not implemented — defaulting to MLA.")
    mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)

def apa_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit=None, include_works_cited_in_limit=False, works_cited_input=None):
    print("APA format not implemented — defaulting to MLA.")
    mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)

def harvard_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit=None, include_works_cited_in_limit=False, works_cited_input=None):
    print("Harvard format not implemented — defaulting to MLA.")
    mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)

def ieee_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit=None, include_works_cited_in_limit=False, works_cited_input=None):
    print("IEEE format not implemented — defaulting to MLA.")
    mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)

# ----------------------
# 공용 실행 로직
# ----------------------
def format_document_by_style(style, input_file, output_file, title, student_name, professor_name, course_name, page_limit=None, include_works_cited_in_limit=False, works_cited_input=None):
    style = (style or "mla").strip().lower()

    if style.startswith('m'):
        mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)
    elif style.startswith('c'):
        chicago_format(input_file, output_file, title, student_name or None, professor_name or None, course_name or None, page_limit, include_works_cited_in_limit, works_cited_input)
    elif style.startswith('a'):
        apa_format(input_file, output_file, title, student_name or None, professor_name or None, course_name or None, page_limit, include_works_cited_in_limit, works_cited_input)
    elif style.startswith('h'):
        harvard_format(input_file, output_file, title, student_name or None, professor_name or None, course_name or None, page_limit, include_works_cited_in_limit, works_cited_input)
    elif style.startswith('i'):
        ieee_format(input_file, output_file, title, student_name or None, professor_name or None, course_name or None, page_limit, include_works_cited_in_limit, works_cited_input)
    else:
        print("Unknown style specified — defaulting to MLA.")
        mla_format(input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)


def run_cli():
    print("=== Word Formatter ===")
    style = input("Choose style (MLA/Chicago/APA/Harvard/IEEE) [MLA]: ").strip().lower() or "mla"
    title = input("Enter your essay title: ")
    student_name = input("Enter your name (optional for Chicago): ")
    professor_title_choice = input("Choose professor title (doctor/professor) [professor]: ").strip().lower() or "professor"
    professor_title = "Doctor" if professor_title_choice.startswith('d') else "Professor"
    professor_only_name = input("Enter professor's name (without title): ").strip()
    professor_name = f"{professor_title} {professor_only_name}" if professor_only_name else ""
    course_name = input("Enter course name: ")
    page_limit_input = input("Enter page limit (or press Enter to skip): ")
    page_limit = int(page_limit_input) if page_limit_input.strip() else None
    include_works_cited_in_limit = False
    if page_limit is not None:
        include_input = input("Include 'Works Cited' page in page limit? (y/N): ").strip().lower()
        include_works_cited_in_limit = include_input.startswith('y')
    works_cited_input = input("Enter works cited entries (one per line, or 'author: title' to auto-fetch, or press Enter to skip): ")
    input_file = input("Enter path to your Word (.docx) file: ")

    # 출력 파일 이름
    output_file = os.path.splitext(input_file)[0] + "_formatted.docx"

    format_document_by_style(style, input_file, output_file, title, student_name, professor_name, course_name, page_limit, include_works_cited_in_limit, works_cited_input)


class WordFormatterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word Formatter")
        self.root.geometry("760x760")
        self.root.minsize(700, 650)

        self.style_var = tk.StringVar(value="MLA")
        self.title_var = tk.StringVar()
        self.student_var = tk.StringVar()
        self.prof_title_var = tk.StringVar(value="professor")
        self.prof_name_var = tk.StringVar()
        self.course_var = tk.StringVar()
        self.page_limit_var = tk.StringVar()
        self.include_wc_var = tk.BooleanVar(value=False)
        self.input_file_var = tk.StringVar()

        self._build_ui()

    def _build_ui(self):
        frame = ttk.Frame(self.root, padding=16)
        frame.pack(fill=tk.BOTH, expand=True)

        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(9, weight=1)

        ttk.Label(frame, text="Style").grid(row=0, column=0, sticky="w", pady=4)
        ttk.Combobox(frame, textvariable=self.style_var, values=["MLA", "Chicago", "APA", "Harvard", "IEEE"], state="readonly").grid(row=0, column=1, sticky="ew", pady=4)

        ttk.Label(frame, text="Essay title").grid(row=1, column=0, sticky="w", pady=4)
        ttk.Entry(frame, textvariable=self.title_var).grid(row=1, column=1, sticky="ew", pady=4)

        ttk.Label(frame, text="Student name").grid(row=2, column=0, sticky="w", pady=4)
        ttk.Entry(frame, textvariable=self.student_var).grid(row=2, column=1, sticky="ew", pady=4)

        ttk.Label(frame, text="Professor title").grid(row=3, column=0, sticky="w", pady=4)
        ttk.Combobox(frame, textvariable=self.prof_title_var, values=["doctor", "professor"], state="readonly").grid(row=3, column=1, sticky="ew", pady=4)

        ttk.Label(frame, text="Professor name (without title)").grid(row=4, column=0, sticky="w", pady=4)
        ttk.Entry(frame, textvariable=self.prof_name_var).grid(row=4, column=1, sticky="ew", pady=4)

        ttk.Label(frame, text="Course name").grid(row=5, column=0, sticky="w", pady=4)
        ttk.Entry(frame, textvariable=self.course_var).grid(row=5, column=1, sticky="ew", pady=4)

        ttk.Label(frame, text="Page limit (optional)").grid(row=6, column=0, sticky="w", pady=4)
        ttk.Entry(frame, textvariable=self.page_limit_var).grid(row=6, column=1, sticky="ew", pady=4)

        ttk.Checkbutton(frame, text="Include Works Cited in page limit", variable=self.include_wc_var).grid(row=7, column=1, sticky="w", pady=4)

        ttk.Label(frame, text="Input .docx file").grid(row=8, column=0, sticky="w", pady=4)
        file_row = ttk.Frame(frame)
        file_row.grid(row=8, column=1, sticky="ew", pady=4)
        file_row.columnconfigure(0, weight=1)
        ttk.Entry(file_row, textvariable=self.input_file_var).grid(row=0, column=0, sticky="ew")
        ttk.Button(file_row, text="Browse", command=self._browse_file).grid(row=0, column=1, padx=(8, 0))

        ttk.Label(frame, text="Works cited entries (one per line)").grid(row=9, column=0, sticky="nw", pady=4)
        self.works_cited_text = scrolledtext.ScrolledText(frame, wrap=tk.WORD, height=10)
        self.works_cited_text.grid(row=9, column=1, sticky="nsew", pady=4)

        action_row = ttk.Frame(frame)
        action_row.grid(row=10, column=1, sticky="e", pady=(12, 0))
        self.run_button = ttk.Button(action_row, text="Format Document", command=self._start_formatting)
        self.run_button.pack(side=tk.RIGHT)

    def _browse_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Word document",
            filetypes=[("Word Document", "*.docx")]
        )
        if file_path:
            self.input_file_var.set(file_path)

    def _start_formatting(self):
        self.run_button.config(state=tk.DISABLED)
        threading.Thread(target=self._format_document, daemon=True).start()

    def _format_document(self):
        try:
            style = self.style_var.get().strip() or "MLA"
            title = self.title_var.get().strip()
            student_name = self.student_var.get().strip()
            prof_title_choice = self.prof_title_var.get().strip().lower() or "professor"
            prof_title = "Doctor" if prof_title_choice.startswith('d') else "Professor"
            prof_only_name = self.prof_name_var.get().strip()
            professor_name = f"{prof_title} {prof_only_name}" if prof_only_name else ""
            course_name = self.course_var.get().strip()
            input_file = self.input_file_var.get().strip()
            works_cited_input = self.works_cited_text.get("1.0", tk.END).strip()

            if not title:
                raise ValueError("Essay title is required.")
            if not input_file:
                raise ValueError("Input .docx file is required.")
            if not os.path.exists(input_file):
                raise ValueError("Input file does not exist.")
            if not input_file.lower().endswith(".docx"):
                raise ValueError("Please select a .docx file.")

            page_limit_text = self.page_limit_var.get().strip()
            page_limit = int(page_limit_text) if page_limit_text else None

            output_file = os.path.splitext(input_file)[0] + "_formatted.docx"

            format_document_by_style(
                style,
                input_file,
                output_file,
                title,
                student_name,
                professor_name,
                course_name,
                page_limit,
                self.include_wc_var.get(),
                works_cited_input
            )

            self.root.after(0, lambda: messagebox.showinfo("Success", f"Formatted document saved:\n{output_file}"))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", str(e)))
        finally:
            self.root.after(0, lambda: self.run_button.config(state=tk.NORMAL))


def run_gui():
    root = tk.Tk()
    WordFormatterApp(root)
    root.mainloop()


def main():
    parser = argparse.ArgumentParser(description="Word Formatter application")
    parser.add_argument("--cli", action="store_true", help="Run in CLI mode instead of GUI mode")
    args = parser.parse_args()

    if args.cli:
        run_cli()
    else:
        run_gui()

if __name__ == "__main__":
    main()